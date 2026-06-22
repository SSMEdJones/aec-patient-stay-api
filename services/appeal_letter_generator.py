"""
Appeal Letter Generator

Generates completed Medicare appeal letters by:
1. Fetching patient data from Epic FHIR
2. Generating MidnightReason justifications using LLM
3. Filling in the docx template with all data
"""
import re
import os
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict
from dataclasses import dataclass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.midnight_reason_generator import (
    MidnightReasonGenerator, 
    PatientStayData, 
    MidnightReasonOutput
)


@dataclass
class AppealLetterData:
    """All data needed to fill the appeal letter template."""
    # Patient info
    member_name: str = ""
    dob: str = ""
    age: str = ""
    gender: str = ""
    member_id: str = ""
    medical_history: str = ""  # PMH abbreviations
    hook: str = ""  # Opening statement hook (chief complaint)
    
    # Payer info
    payer_name: str = ""
    street_address: str = ""
    city: str = ""
    state: str = ""
    zip_code: str = ""
    
    # Case info
    reference_number: str = ""
    dos: str = ""  # Date of Service
    place_of_service: str = ""  # Hospital/facility name
    
    # Generated content
    patient_background: str = ""
    midnight_reason_1: str = ""
    midnight_reason_2: str = ""
    closing_summary: str = ""
    
    # Ministry/Facility info (for dynamic header)
    ministry_name: str = ""
    ministry_address: str = ""
    ministry_city: str = ""
    ministry_state: str = ""
    ministry_zip: str = ""
    ministry_phone: str = ""


class AppealLetterGenerator:
    """
    Generates completed Medicare appeal letters from Epic patient data.
    
    Usage:
        generator = AppealLetterGenerator()
        output_path = generator.generate(
            patient_id="erXuFYUfucBZaryVksYEcMg3",
            payer_info={...},
            output_dir="output"
        )
    """
    
    def __init__(self, template_path: str = "appeal_templates/Template.docx"):
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        self.reason_generator = MidnightReasonGenerator()
    
    def _replace_placeholder(self, paragraph, placeholder: str, value: str):
        """Replace a placeholder in a paragraph, handling split runs."""
        full_text = paragraph.text
        if placeholder not in full_text:
            return False
        
        # Simple case: placeholder is in a single run
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return True
        
        # Complex case: placeholder spans multiple runs
        # Rebuild the paragraph text
        new_text = full_text.replace(placeholder, value)
        
        # Clear all runs and set text in first run
        if paragraph.runs:
            # Preserve formatting from first run
            first_run = paragraph.runs[0]
            for run in paragraph.runs[1:]:
                run.text = ""
            first_run.text = new_text
            return True
        
        return False
    
    def _replace_in_shapes(self, element, replacements: dict):
        """Replace placeholders in text boxes and shapes within an element.
        
        This handles text inside shapes that were converted from EMF/WMF graphics.
        Shape text is stored in <w:txbxContent> elements in the XML.
        Handles placeholders split across multiple runs while preserving line breaks.
        """
        from docx.oxml.ns import qn
        from lxml import etree
        
        # Find all text box content elements
        for txbx in element.iter(qn('w:txbxContent')):
            # Find paragraphs that are direct children of this text box
            for para_elem in txbx.findall(qn('w:p')):
                # Collect runs and track which have line breaks after them
                runs_data = []  # List of (run_elem, text, has_br_after)
                run_elems = para_elem.findall(qn('w:r'))
                
                for run_elem in run_elems:
                    # Get text from this run
                    text_parts = [t.text or '' for t in run_elem.findall(qn('w:t'))]
                    run_text = ''.join(text_parts)
                    
                    # Check if this run has a line break
                    has_br = run_elem.find(qn('w:br')) is not None
                    
                    runs_data.append((run_elem, run_text, has_br))
                
                if not runs_data:
                    continue
                
                # Build combined text with markers for line breaks
                BR_MARKER = '\x00BR\x00'
                combined_parts = []
                for run_elem, run_text, has_br in runs_data:
                    combined_parts.append(run_text)
                    if has_br:
                        combined_parts.append(BR_MARKER)
                
                full_text = ''.join(combined_parts)
                if not full_text.replace(BR_MARKER, ''):
                    continue
                
                # Replace placeholders
                new_text = full_text
                for placeholder, value in replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, value)
                
                if new_text != full_text:
                    # Collapse consecutive BR markers into single ones
                    while BR_MARKER + BR_MARKER in new_text:
                        new_text = new_text.replace(BR_MARKER + BR_MARKER, BR_MARKER)
                    # Remove leading/trailing BR markers
                    new_text = new_text.strip(BR_MARKER.strip('\x00'))
                    if new_text.startswith(BR_MARKER):
                        new_text = new_text[len(BR_MARKER):]
                    if new_text.endswith(BR_MARKER):
                        new_text = new_text[:-len(BR_MARKER)]
                    
                    print(f"[SHAPE] Replaced: '{full_text.replace(BR_MARKER, '[BR]')}' -> '{new_text.replace(BR_MARKER, '[BR]')}'")
                    
                    # Split by BR markers to get lines
                    lines = new_text.split(BR_MARKER)
                    
                    # Clear existing runs
                    for run_elem, _, _ in runs_data:
                        for t in run_elem.findall(qn('w:t')):
                            t.text = ''
                        # Remove any existing br elements
                        for br in run_elem.findall(qn('w:br')):
                            run_elem.remove(br)
                    
                    # Put text in first run, with line breaks
                    if runs_data:
                        first_run = runs_data[0][0]
                        first_t = first_run.find(qn('w:t'))
                        if first_t is None:
                            first_t = etree.SubElement(first_run, qn('w:t'))
                        
                        # Set first line
                        first_t.text = lines[0] if lines else ''
                        
                        # Add remaining lines with <w:br/> before each
                        for line in lines[1:]:
                            br = etree.SubElement(first_run, qn('w:br'))
                            t = etree.SubElement(first_run, qn('w:t'))
                            t.text = line
    
    def _add_ministry_header(self, doc: Document, data: AppealLetterData):
        """Add ministry letterhead to document header dynamically."""
        # Debug: print what we're working with
        print(f"[HEADER DEBUG] ministry_name: '{data.ministry_name}'")
        print(f"[HEADER DEBUG] ministry_address: '{data.ministry_address}'")
        
        # Access the first section
        section = doc.sections[0]
        
        # Check if template uses different first page header
        if section.different_first_page_header_footer:
            print("[HEADER DEBUG] Template uses different first page header")
            header = section.first_page_header
        else:
            header = section.header
        
        header.is_linked_to_previous = False
        
        # Add a NEW paragraph for ministry info (preserving any existing logo/images)
        ministry_para = header.add_paragraph()
        
        # Right-align for typical letterhead style
        ministry_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Check if ministry info was provided
        if not data.ministry_name:
            # Mark as not found so user knows
            run = ministry_para.add_run("** MINISTRY NOT FOUND **")
            run.bold = True
            run.font.size = Pt(11)
            return
        
        # Add ministry name (bold)
        run = ministry_para.add_run(data.ministry_name)
        run.bold = True
        run.font.size = Pt(11)
        
        # Add address lines
        if data.ministry_address:
            ministry_para.add_run(f'\n{data.ministry_address}')
        
        if data.ministry_city or data.ministry_state or data.ministry_zip:
            city_state_zip = ", ".join(filter(None, [
                data.ministry_city,
                f"{data.ministry_state} {data.ministry_zip}".strip()
            ]))
            ministry_para.add_run(f'\n{city_state_zip}')
        
        if data.ministry_phone:
            ministry_para.add_run(f'\n{data.ministry_phone}')
    
    def _fill_template(self, data: AppealLetterData, output_path: Path) -> Path:
        """Fill the template with the provided data."""
        doc = Document(self.template_path)
        
        # Ministry header is now handled via placeholder replacement in the template
        # Placeholders: [MinistryName], [MinistryAddress], [MinistryCity], [MinistryState], [MinistryZip], [MinistryPhone]
        
        # Mapping of placeholders to values
        # Note: Template has typo "Mednight" instead of "Midnight"
        # [DOSHeader] = full multi-date format for header
        # [DOS] = just the first date for body paragraph
        dos_header = data.dos  # Full format with Observation/Inpatient
        dos_body = data.dos.split(" - ")[0].split("\n")[0] if data.dos else ""  # Just first date
        
        # Clean up MidnightReason - strip trailing punctuation since template adds periods
        mn1_clean = data.midnight_reason_1.rstrip('.,;: ')
        mn2_clean = data.midnight_reason_2.rstrip('.,;: ')
        closing_clean = data.closing_summary.rstrip('.,;: ') if data.closing_summary else ""
        
        # Convert "LAST, FIRST" to "First Last" format for display
        member_name = data.member_name
        if ',' in member_name:
            parts = member_name.split(',', 1)
            last_name = parts[0].strip()
            first_name = parts[1].strip().split()[0] if parts[1].strip() else ''
            member_name = f"{first_name} {last_name}" if first_name else last_name
        
        # Hook is now a complete paragraph - keep capitalization
        # Strip trailing punctuation for consistency
        hook = data.hook
        if hook:
            hook = hook.rstrip('.,;: ')
        
        replacements = {
            "[MemberName]": member_name,
            "[DOB]": data.dob,
            "[Age]": data.age,
            "[Gender]": data.gender,
            "[MemberID]": data.member_id,
            "[MedicalHistory]": data.medical_history,
            "[Hook]": hook,
            "[Complaint]": hook,  # Backwards compatibility
            "[PlaceofService]": data.place_of_service or "Hospital",
            "[Street Address]": data.street_address,
            "[City]": data.city,
            "[State]": data.state,
            "[ZIP]": data.zip_code,
            "[ReferenceNumber]": data.reference_number,
            "[DOSHeader]": dos_header,  # Full format for header
            "[DOS]": dos_body,  # Just date for body
            "[MednightReason1]": mn1_clean,
            "[MednightReason2]": mn2_clean,
            # Also handle correct spelling in case template is fixed
            "[MidnightReason1]": mn1_clean,
            "[MidnightReason2]": mn2_clean,
            "[ClosingSummary]": closing_clean,
            # Ministry/Hospital info for dynamic header
            "[MinistryName]": data.ministry_name,
            "[MinistryAddress]": data.ministry_address,
            "[MinistryCity]": data.ministry_city,
            "[MinistryState]": data.ministry_state,
            "[MinistryZip]": data.ministry_zip,
            "[MinistryPhone]": data.ministry_phone,
        }
        
        # Replace in all paragraphs (document body)
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    self._replace_placeholder(paragraph, placeholder, value)
        
        # Replace in headers (all sections, including first page header)
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header is None:
                    continue
                # Regular paragraphs in header
                for paragraph in header.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            self._replace_placeholder(paragraph, placeholder, value)
                # Text in shapes/text boxes (for converted EMF graphics)
                self._replace_in_shapes(header._element, replacements)
        
        # Also check tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                self._replace_placeholder(paragraph, placeholder, value)
        
        # Post-process: Fix "a" vs "an" grammar for ages
        self._fix_a_an_in_document(doc)
        
        # Add AI disclaimer to document footer (appears at bottom of every page)
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()  # Clear any existing content
        run = footer_para.add_run("This report was generated by the SSM Health AI Portal - Fast Appeal Tool")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        run.font.italic = True
        
        # Save the filled document
        doc.save(output_path)
        return output_path
    
    def _fix_a_an_in_document(self, doc):
        """Fix 'a' vs 'an' grammar for ages in all paragraphs."""
        import re
        
        def fix_text(text):
            def fix_match(match):
                age = match.group(2)
                suffix = match.group(3)
                # Ages that need "an": 8, 11, 18, 80-89
                if age.startswith('8') or age == '11' or age == '18':
                    return f"an {age}{suffix}"
                else:
                    return f"a {age}{suffix}"
            
            text = re.sub(r'\b(a|an)\s+(\d+)(-year-old)', fix_match, text, flags=re.IGNORECASE)
            return text
        
        # Fix in all paragraphs
        for paragraph in doc.paragraphs:
            if 'year-old' in paragraph.text.lower():
                for run in paragraph.runs:
                    if 'year-old' in run.text.lower() or re.search(r'\b(a|an)\s+\d+', run.text, re.IGNORECASE):
                        run.text = fix_text(run.text)
                # If fix didn't work on runs (text split across runs), fix whole paragraph
                if re.search(r'\ba\s+8\d*-year-old', paragraph.text, re.IGNORECASE):
                    full_text = fix_text(paragraph.text)
                    if paragraph.runs:
                        paragraph.runs[0].text = full_text
                        for run in paragraph.runs[1:]:
                            run.text = ""
        
        # Also fix in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'year-old' in paragraph.text.lower():
                            for run in paragraph.runs:
                                run.text = fix_text(run.text)
    
    def _format_medical_history(self, conditions: list) -> str:
        """Format conditions list into medical history with abbreviations."""
        # Common abbreviation mappings
        abbreviations = {
            "hypertension": "HTN",
            "essential hypertension": "HTN",
            "diabetes mellitus": "DM",
            "diabetes": "DM",
            "type 2 diabetes": "DM type 2",
            "chronic obstructive pulmonary disease": "COPD",
            "congestive heart failure": "CHF",
            "heart failure": "CHF",
            "coronary artery disease": "CAD",
            "atrial fibrillation": "AFib",
            "paroxysmal atrial fibrillation": "PAF",
            "hyperlipidemia": "HLD",
            "high cholesterol": "HLD",
            "hypercholesterolemia": "HLD",
            "chronic kidney disease": "CKD",
            "end stage renal disease": "ESRD",
            "peripheral vascular disease": "PVD",
            "cerebrovascular accident": "CVA",
            "transient ischemic attack": "TIA",
            "gastroesophageal reflux disease": "GERD",
            "deep vein thrombosis": "DVT",
            "pulmonary embolism": "PE",
            "hypothyroidism": "hypothyroidism",
            "hyperthyroidism": "hyperthyroidism",
            "polycystic ovarian syndrome": "PCOS",
        }
        
        formatted = []
        for cond in conditions:
            cond_lower = cond.lower()
            # Check for abbreviation match
            abbreviated = False
            for full, abbr in abbreviations.items():
                if full in cond_lower:
                    formatted.append(abbr)
                    abbreviated = True
                    break
            if not abbreviated:
                formatted.append(cond)
        
        # Remove duplicates while preserving order
        seen = set()
        unique = []
        for item in formatted:
            if item.lower() not in seen:
                seen.add(item.lower())
                unique.append(item)
        
        # Format with proper grammar: "A, B, and C"
        if len(unique) == 0:
            return "See clinical documentation"
        elif len(unique) == 1:
            return unique[0]
        elif len(unique) == 2:
            return f"{unique[0]} and {unique[1]}"
        else:
            return ", ".join(unique[:-1]) + ", and " + unique[-1]
    
    def generate(
        self,
        patient_id: str,
        payer_info: Optional[Dict[str, str]] = None,
        reference_number: str = "",
        member_id: str = "",
        place_of_service: str = "",
        output_dir: str = "output",
        output_filename: Optional[str] = None
    ) -> Path:
        """
        Generate a complete appeal letter for a patient.
        
        Args:
            patient_id: Epic FHIR Patient ID
            payer_info: Dict with payer details (name, street_address, city, state, zip)
            reference_number: Appeal reference number
            member_id: Insurance member ID (if different from MRN)
            place_of_service: Hospital/facility name
            output_dir: Directory to save the generated letter
            output_filename: Optional custom filename
            
        Returns:
            Path to the generated docx file
        """
        # Default payer info (can be customized)
        payer = payer_info or {
            "name": "Medicare Advantage Plan",
            "street_address": "PO Box 0000",
            "city": "City",
            "state": "ST",
            "zip": "00000"
        }
        
        # Fetch patient data and generate MidnightReason
        print(f"Fetching patient data for: {patient_id}")
        patient_data = self.reason_generator.epic_fetcher.fetch_patient_stay_data(patient_id)
        
        print("Generating MidnightReason justifications...")
        reason_output = self.reason_generator.generate_from_data(patient_data)
        
        # Format dates
        dob_formatted = patient_data.dob
        if dob_formatted:
            try:
                dob_dt = datetime.strptime(dob_formatted[:10], "%Y-%m-%d")
                dob_formatted = dob_dt.strftime("%m/%d/%Y")
            except:
                pass
        
        # Format DOS with observation/inpatient status if available
        dos_formatted = ""
        obs_date = getattr(patient_data, 'observation_date', '') or ''
        inp_date = getattr(patient_data, 'inpatient_date', '') or ''
        adm_date = patient_data.admission_date or ''
        
        def fmt_date(d):
            if not d:
                return ""
            try:
                dt = datetime.strptime(d[:10], "%Y-%m-%d")
                return dt.strftime("%m/%d/%Y")
            except:
                return d
        
        if obs_date and inp_date:
            # Both dates - show transition
            dos_formatted = f"{fmt_date(obs_date)} (Observation),\n  {fmt_date(inp_date)} (Inpatient, current)"
        elif obs_date:
            dos_formatted = f"{fmt_date(obs_date)} (Observation)"
        elif inp_date:
            dos_formatted = f"{fmt_date(inp_date)} (Inpatient)"
        elif adm_date:
            dos_formatted = f"{fmt_date(adm_date)} to current"
        
        # Format medical history from conditions (abbreviate common terms)
        medical_history = self._format_medical_history(patient_data.conditions)
        
        # Format gender (handle both full names and single letters)
        gender = patient_data.gender.lower() if patient_data.gender else ""
        gender_display = "male" if gender in ("male", "m") else "female" if gender in ("female", "f") else gender
        
        # Format hook (chief complaint or first condition)
        hook = patient_data.chief_complaint
        if not hook and patient_data.conditions:
            hook = patient_data.conditions[0].lower()
        hook = hook or "evaluation and management"
        
        # Generate random member ID and reference number if not provided
        import random
        random_member_id = f"{random.randint(100000000, 999999999)}"
        random_ref_num = f"A{random.randint(100000000, 999999999)}"
        
        # Use authorization_number from patient data if available, else fall back to parameter or random
        effective_ref_num = patient_data.authorization_number or reference_number or random_ref_num
        # Use insurance_id from patient data if available, else fall back to parameter or random
        effective_member_id = patient_data.insurance_id or member_id or random_member_id
        
        # Log what we're using
        print(f"[DEBUG] patient_data.authorization_number = '{patient_data.authorization_number}'")
        print(f"[DEBUG] reference_number param = '{reference_number}'")
        print(f"[DEBUG] effective_ref_num = '{effective_ref_num}'")
        print(f"[DEBUG] patient_data.insurance_id = '{patient_data.insurance_id}'")
        print(f"[DEBUG] effective_member_id = '{effective_member_id}'")
        
        # Prepare letter data
        letter_data = AppealLetterData(
            member_name=patient_data.patient_name,
            dob=dob_formatted,
            age=str(patient_data.age) if patient_data.age else "",
            gender=gender_display,
            member_id=effective_member_id,
            medical_history=medical_history,
            hook=hook,
            place_of_service=place_of_service or "Hospital",
            street_address=payer.get("street_address", ""),
            city=payer.get("city", ""),
            state=payer.get("state", ""),
            zip_code=payer.get("zip", ""),
            reference_number=effective_ref_num,
            dos=dos_formatted,
            patient_background=reason_output.patient_background,
            midnight_reason_1=reason_output.midnight_reason_1,
            midnight_reason_2=reason_output.midnight_reason_2,
            closing_summary=reason_output.closing_summary,
        )
        
        # Create output directory
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # Generate filename
        if output_filename:
            filename = output_filename
        else:
            safe_name = re.sub(r'[^\w\s-]', '', patient_data.patient_name).strip()
            safe_name = re.sub(r'[\s]+', '_', safe_name)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Appeal_{safe_name}_{timestamp}.docx"
        
        output_file = output_path / filename
        
        # Fill template and save
        print(f"Generating appeal letter: {output_file}")
        self._fill_template(letter_data, output_file)
        
        print(f"Appeal letter saved to: {output_file}")
        return output_file
    
    def generate_from_data(
        self,
        patient_data: PatientStayData,
        reason_output: MidnightReasonOutput,
        payer_info: Optional[Dict[str, str]] = None,
        reference_number: str = "",
        member_id: str = "",
        place_of_service: str = "",
        output_dir: str = "output",
        output_filename: Optional[str] = None
    ) -> Path:
        """
        Generate appeal letter from pre-fetched data.
        
        Useful when you've already fetched data or want to use mock data.
        """
        payer = payer_info or {
            "name": "Medicare Advantage Plan",
            "street_address": "PO Box 0000",
            "city": "City",
            "state": "ST",
            "zip": "00000"
        }
        
        # Format dates
        dob_formatted = patient_data.dob
        if dob_formatted:
            try:
                dob_dt = datetime.strptime(dob_formatted[:10], "%Y-%m-%d")
                dob_formatted = dob_dt.strftime("%m/%d/%Y")
            except:
                pass
        
        # Format DOS with observation/inpatient status if available
        import random
        dos_formatted = ""
        obs_date = getattr(patient_data, 'observation_date', '') or ''
        inp_date = getattr(patient_data, 'inpatient_date', '') or ''
        adm_date = patient_data.admission_date or ''
        
        def fmt_date(d):
            if not d:
                return ""
            try:
                dt = datetime.strptime(d[:10], "%Y-%m-%d")
                return dt.strftime("%m/%d/%Y")
            except:
                return d
        
        if obs_date and inp_date:
            # Both dates - show transition
            dos_formatted = f"{fmt_date(obs_date)} (Observation),\n  {fmt_date(inp_date)} (Inpatient, current)"
        elif obs_date:
            dos_formatted = f"{fmt_date(obs_date)} (Observation)"
        elif inp_date:
            dos_formatted = f"{fmt_date(inp_date)} (Inpatient)"
        elif adm_date:
            dos_formatted = f"{fmt_date(adm_date)} to current"
        
        # Format medical history
        medical_history = self._format_medical_history(patient_data.conditions)
        
        # Format gender (handle both full names and single letters)
        gender = patient_data.gender.lower() if patient_data.gender else ""
        gender_display = "male" if gender in ("male", "m") else "female" if gender in ("female", "f") else gender
        
        # Format complaint
        hook = patient_data.chief_complaint
        if not hook and patient_data.conditions:
            hook = patient_data.conditions[0].lower()
        hook = hook or "evaluation and management"
        
        # Generate random member ID and reference number
        random_member_id = f"{random.randint(100000000, 999999999)}"
        random_ref_num = f"A{random.randint(100000000, 999999999)}"
        
        # Use authorization_number from patient data if available, else fall back to parameter or random
        effective_ref_num = patient_data.authorization_number or reference_number or random_ref_num
        # Use insurance_id from patient data if available, else fall back to parameter or random
        effective_member_id = patient_data.insurance_id or member_id or random_member_id
        
        # Log what we're using
        print(f"[DEBUG] patient_data.authorization_number = '{patient_data.authorization_number}'")
        print(f"[DEBUG] effective_ref_num = '{effective_ref_num}'")
        print(f"[DEBUG] patient_data.insurance_id = '{patient_data.insurance_id}'")
        print(f"[DEBUG] effective_member_id = '{effective_member_id}'")
        
        letter_data = AppealLetterData(
            member_name=patient_data.patient_name,
            dob=dob_formatted,
            age=str(patient_data.age) if patient_data.age else "",
            gender=gender_display,
            member_id=effective_member_id,
            medical_history=medical_history,
            hook=hook,
            place_of_service=place_of_service or "Hospital",
            street_address=payer.get("street_address", ""),
            city=payer.get("city", ""),
            state=payer.get("state", ""),
            zip_code=payer.get("zip", ""),
            reference_number=effective_ref_num,
            dos=dos_formatted,
            patient_background=reason_output.patient_background,
            midnight_reason_1=reason_output.midnight_reason_1,
            midnight_reason_2=reason_output.midnight_reason_2,
            closing_summary=reason_output.closing_summary,
        )
        
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        if output_filename:
            filename = output_filename
        else:
            safe_name = re.sub(r'[^\w\s-]', '', patient_data.patient_name).strip()
            safe_name = re.sub(r'[\s]+', '_', safe_name)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Appeal_{safe_name}_{timestamp}.docx"
        
        output_file = output_path / filename
        self._fill_template(letter_data, output_file)
        
        return output_file


if __name__ == "__main__":
    import sys
    
    # Test with Epic patient or mock data
    generator = AppealLetterGenerator()
    
    if len(sys.argv) > 1:
        patient_id = sys.argv[1]
    else:
        patient_id = "erXuFYUfucBZaryVksYEcMg3"  # Camila Lopez
    
    # Example payer info (customize as needed)
    payer_info = {
        "name": "AARP Medicare Advantage Essentials UHC ST-3",
        "street_address": "PO Box 6106, MS CA120-0360",
        "city": "Cypress",
        "state": "CA",
        "zip": "90630-0016"
    }
    
    output = generator.generate(
        patient_id=patient_id,
        payer_info=payer_info,
        reference_number="A998064317",
        member_id="969345931"
    )
    
    print(f"\nGenerated: {output}")
