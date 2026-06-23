from docx import Document
from docx.oxml.ns import qn

doc = Document('appeal_templates/Template.docx')
print('Sections:', len(doc.sections))
section = doc.sections[0]
print('Different first page header:', section.different_first_page_header_footer)

# Check first page header
header = section.first_page_header
print('\n=== First Page Header ===')
print('Paragraphs:', len(header.paragraphs))
for i, p in enumerate(header.paragraphs):
    txt = p.text[:100] + '...' if len(p.text) > 100 else p.text
    print(f'  Para {i}: "{txt}"')
    # Check for images (drawings) in runs
    for run in p.runs:
        drawings = run._r.findall('.//'+qn('w:drawing'))
        if drawings:
            print(f'    [Has {len(drawings)} image(s)]')

# Check regular header
header2 = section.header
print('\n=== Regular Header ===')
print('Paragraphs:', len(header2.paragraphs))
for i, p in enumerate(header2.paragraphs):
    txt = p.text[:100] + '...' if len(p.text) > 100 else p.text
    print(f'  Para {i}: "{txt}"')
    for run in p.runs:
        drawings = run._r.findall('.//'+qn('w:drawing'))
        if drawings:
            print(f'    [Has {len(drawings)} image(s)]')
